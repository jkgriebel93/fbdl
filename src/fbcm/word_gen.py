import json
import math
import os
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor, Twips
from PIL import Image, ImageDraw, ImageFont

from .constants import POSITION_STATS
from .models import ColorScheme, ProspectDataSoup


def get_primary_position(position: str) -> str:
    """Extract the primary position from multi-position strings like 'DL/EDGE'."""
    if not position:
        return ""
    return position.split("/")[0].strip().upper()


class SchoolColors:
    def __init__(self, colors_file: str):
        with open(colors_file, "r") as infile:
            self.color_data = self._normalize_color_data(data=json.load(infile))

    def _normalize_color_data(self, data: Dict) -> Dict:
        normalized = {}
        for division in ["FBS", "FCS"]:
            if division not in data:
                continue
            for conference, schools in data[division].items():
                for school, colors in schools.items():
                    normalized[school.lower()] = colors

        return normalized

    def get_school_colors(self, school: str) -> ColorScheme:
        normalized = school.lower().strip()
        colors = ColorScheme(**self.color_data[normalized])
        colors.dark = self.darken_color(colors.primary, 0.3)
        colors.medium = self.blend_colors(colors.light,
                                          colors.primary,
                                          0.20)

        colors.primary_rgb = self.hex_to_rgb(colors.primary)
        colors.light_rgb = self.hex_to_rgb(colors.light)
        return colors

    def blend_colors(self, color1: str, color2: str, ratio: float = 0.5) -> str:
        """Blend two hex colors."""
        c1 = color1.lstrip("#")
        c2 = color2.lstrip("#")
        r1, g1, b1 = int(c1[0:2], 16), int(c1[2:4], 16), int(c1[4:6], 16)
        r2, g2, b2 = int(c2[0:2], 16), int(c2[2:4], 16), int(c2[4:6], 16)
        r = int(r1 + (r2 - r1) * ratio)
        g = int(g1 + (g2 - g1) * ratio)
        b = int(b1 + (b2 - b1) * ratio)
        return f"{r:02x}{g:02x}{b:02x}"

    def darken_color(self, hex_color: str, factor: float = 0.3) -> str:
        """Darken a hex color."""
        c = hex_color.lstrip("#")
        r = int(int(c[0:2], 16) * (1 - factor))
        g = int(int(c[2:4], 16) * (1 - factor))
        b = int(int(c[4:6], 16) * (1 - factor))
        return f"{r:02x}{g:02x}{b:02x}"

    def hex_to_rgb(self, hex_color: str) -> RGBColor:
        """Convert hex to RGBColor."""
        c = hex_color.lstrip("#")
        return RGBColor(int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16))


def create_rating_ring(
    rating: float,
    primary_color: str,
    light_color: str,
    size: int = 120,
    output_path: str = None,
) -> str:
    """Create a circular progress ring image."""
    ring_width = 12
    img = Image.new("RGBA", (size, size), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)

    pc = primary_color.lstrip("#")
    lc = light_color.lstrip("#")
    primary_rgb = tuple(int(pc[i : i + 2], 16) for i in (0, 2, 4))
    light_rgb = tuple(int(lc[i : i + 2], 16) for i in (0, 2, 4))

    padding = 4
    center = size // 2
    outer_radius = size // 2 - padding
    inner_radius = outer_radius - ring_width
    bbox = [padding, padding, size - padding, size - padding]

    draw.arc(bbox, 0, 360, fill=light_rgb, width=ring_width)

    if rating > 0:
        start_angle = -90
        sweep_angle = (rating / 100) * 360
        end_angle = start_angle + sweep_angle
        draw.arc(bbox, start_angle, end_angle, fill=primary_rgb, width=ring_width)

        cap_radius = ring_width // 2
        start_x, start_y = center, padding + ring_width // 2
        draw.ellipse(
            [
                start_x - cap_radius,
                start_y - cap_radius,
                start_x + cap_radius,
                start_y + cap_radius,
            ],
            fill=primary_rgb,
        )

        if rating < 100:
            end_angle_rad = math.radians(end_angle)
            arc_center_radius = outer_radius - ring_width // 2
            end_x = center + arc_center_radius * math.cos(end_angle_rad)
            end_y = center + arc_center_radius * math.sin(end_angle_rad)
            draw.ellipse(
                [
                    end_x - cap_radius,
                    end_y - cap_radius,
                    end_x + cap_radius,
                    end_y + cap_radius,
                ],
                fill=primary_rgb,
            )

    center_radius = inner_radius - 4
    draw.ellipse(
        [
            center - center_radius,
            center - center_radius,
            center + center_radius,
            center + center_radius,
        ],
        fill=primary_rgb,
    )

    try:
        font = ImageFont.truetype(
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", size // 4
        )
    except:
        font = ImageFont.load_default()

    text = str(int(rating)) if rating == int(rating) else f"{rating:.1f}"
    text_bbox = draw.textbbox((0, 0), text, font=font)
    text_width = text_bbox[2] - text_bbox[0]
    text_height = text_bbox[3] - text_bbox[1]
    draw.text(
        (center - text_width // 2, center - text_height // 2 - 2),
        text,
        fill=(255, 255, 255),
        font=font,
    )

    if output_path:
        img.save(output_path, "PNG")

    return output_path


# ─── DOCX HELPERS ────────────────────────────────────────────────────────────


def set_cell_shading(cell, hex_color: str):
    """Set cell background color."""
    hex_color = hex_color.lstrip("#").upper()
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shading
    for shd in tcPr.findall(qn("w:shd")):
        tcPr.remove(shd)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color)
    tcPr.append(shading)


def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Set cell margins in twips (1/20 of a point, 1440 twips = 1 inch)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing margins
    for tcMar in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(tcMar)

    tcMar = OxmlElement("w:tcMar")
    for margin_name, margin_value in [
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
    ]:
        margin = OxmlElement(f"w:{margin_name}")
        margin.set(qn("w:w"), str(margin_value))
        margin.set(qn("w:type"), "dxa")
        tcMar.append(margin)
    tcPr.append(tcMar)


def remove_cell_borders(cell):
    """Remove all borders from a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for tcBorders in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(tcBorders)

    tcBorders = OxmlElement("w:tcBorders")
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "nil")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_left_border(cell, hex_color: str, size: int = 24):
    """Add left border to cell."""
    hex_color = hex_color.lstrip("#").upper()
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for tcBorders in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(tcBorders)

    tcBorders = OxmlElement("w:tcBorders")
    for border_name in ["top", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "nil")
        tcBorders.append(border)

    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:color"), hex_color)
    tcBorders.append(left)
    tcPr.append(tcBorders)


def skill_bar(pct: int) -> str:
    """Generate ASCII skill bar."""
    filled = round(pct / 10)
    return "█" * filled + "░" * (10 - filled)


# ─── DOCUMENT GENERATION ─────────────────────────────────────────────────────


class WordDocGenerator:
    def __init__(self,
                 prospect: ProspectDataSoup,
                 output_path: str,
                 ring_image_base_dir: str,
                 colors_path: str):
        self.prospect = prospect
        self.output_path = output_path
        self.ring_img_base_dir = ring_image_base_dir
        self.ring_img_path = None
        self.color_handler = SchoolColors(colors_file=colors_path)
        self.colors = self.color_handler.get_school_colors(self.prospect.basic_info.college)

        self.document = Document()

    def _set_margins(self):
        section = self.document.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(0.75)

    def _gen_rating_ring(self, size: int = 120):
        ring_width = 12
        img = Image.new("RGBA", (size, size), (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)

        padding = 4
        center = size // 2
        outer_radius = size // 2 - padding
        inner_radius = outer_radius - ring_width
        bbox = [padding, padding, size - padding, size - padding]

        draw.arc(bbox,
                 start=0,
                 end=360,
                 fill=self.colors.light_rgb,
                 width=ring_width)

        rating = self.prospect.ratings.overall_rating

        if rating > 0:
            start_angle = -90
            sweep_angle = (rating / 100) * 360
            end_angle = start_angle + sweep_angle
            draw.arc(bbox,
                     start_angle,
                     end_angle,
                     fill=self.colors.primary_rgb,
                     width=ring_width)

            cap_radius = ring_width // 2
            start_x, start_y = center, padding + ring_width // 2
            draw.ellipse(
                [
                    start_x - cap_radius,
                    start_y - cap_radius,
                    start_x + cap_radius,
                    start_y + cap_radius,
                ],
                fill=self.colors.primary_rgb,
            )

            if rating < 100:
                end_angle_rad = math.radians(end_angle)
                arc_center_radius = outer_radius - ring_width // 2
                end_x = center + arc_center_radius * math.cos(end_angle_rad)
                end_y = center + arc_center_radius * math.sin(end_angle_rad)
                draw.ellipse(
                    [
                        end_x - cap_radius,
                        end_y - cap_radius,
                        end_x + cap_radius,
                        end_y + cap_radius,
                    ],
                    fill=self.colors.primary_rgb,
                )

        center_radius = inner_radius - 4
        draw.ellipse(
            [
                center - center_radius,
                center - center_radius,
                center + center_radius,
                center + center_radius,
            ],
            fill=self.colors.primary_rgb,
        )

        try:
            # TODO: Parameterize the fontpath here
            font = ImageFont.truetype(
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", size // 4
            )
        except:
            font = ImageFont.load_default()

        text = str(int(rating)) if rating == int(rating) else f"{rating:.1f}"
        text_bbox = draw.textbbox((0, 0), text, font=font)
        text_width = text_bbox[2] - text_bbox[0]
        text_height = text_bbox[3] - text_bbox[1]

        draw.text(
            (center - text_width // 2, center - text_height // 2 - 2),
            text,
            fill=(255, 255, 255),
            font=font,
        )

        self.ring_img_path = f"{self.ring_img_base_dir}_{self.prospect.basic_info.full_name}_ring.png"
        img.save(self.ring_img_path, "PNG")
        return self.ring_img_path

    def _gen_header_table(self):
        # TODO: Split this into smaller methods
        header_table = self.document.add_table(rows=1, cols=3)
        header_table.autofit = False

        photo_cell = header_table.cell(0, 0)
        photo_cell.width = Inches(1.5)
        remove_cell_borders(photo_cell)
        photo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        photo_para = photo_cell.paragraphs[0]
        photo_para.add_run().add_picture(str(self.prospect.basic_info.photo_path), width=Inches(1.3))

        name_cell = header_table.cell(0, 1)
        name_cell.width = Inches(4.0)
        remove_cell_borders(name_cell)
        name_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        name_para = name_cell.paragraphs[0]
        name_para.paragraph_format.space_after = Pt(0)
        run = name_para.add_run(self.prospect.basic_info.full_name)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = self.colors.primary_rgb

        info_para = name_cell.add_paragraph()
        info_para.paragraph_format.space_before = Pt(2)
        info_para.paragraph_format.space_after = Pt(4)
        run = info_para.add_run(
            f"{self.prospect.basic_info.position}  •  "
            f"{self.prospect.basic_info.college.title()}  •  "
            f"{self.prospect.basic_info.play_style}"
        )
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        measurables_para = name_cell.add_paragraph()
        run = measurables_para.add_run(
            f"{self.prospect.basic_info.height}  •  "
            f"{self.prospect.basic_info.weight} lbs  •  "
            f"{self.prospect.basic_info.forty}s  •  "
            f"{self.prospect.basic_info.hometown.title()}  •  "
            f"{self.prospect.basic_info.class_.title()}"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        ring_cell = header_table.cell(0, 2)
        ring_cell.width = Inches(1.25)
        remove_cell_borders(ring_cell)
        ring_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        ring_para = ring_cell.paragraphs[0]
        ring_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        ring_image_path = self._gen_rating_ring()
        ring_para.add_run().add_picture(ring_image_path, width=Inches(0.95))

        label_para = ring_cell.add_paragraph()
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_para.paragraph_format.space_before = Pt(2)
        run = label_para.add_run("PROSPECT RATING")
        run.font.size = Pt(6.5)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def _gen_rankings_bar(self):
        rankings_table = self.document.add_table(rows=1, cols=5)
        rankings_table.autofit = False

        rankings_data = [
            ("DRAFT BUZZ", str(self.prospect.ratings.overall_rank), "OVERALL", self.colors.primary),
            ("DRAFT BUZZ", str(self.prospect.ratings.position_rank), "POSITION", self.colors.dark),
            ("DRAFT", str(self.prospect.ratings.draft_projection), "PROJECTION", self.colors.primary),
            ("CONSENSUS", str(self.prospect.ratings.avg_overall_rank), "OVERALL", self.colors.dark),
            ("CONSENSUS", str(self.prospect.ratings.avg_position_rank), "POSITION", self.colors.primary),
        ]

        for i, (source, value, label, bg_color) in enumerate(rankings_data):
            cell = rankings_table.cell(0, i)
            cell.width = Inches(1.35)
            remove_cell_borders(cell)
            set_cell_shading(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=80, left=80, right=80)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # Source label
            p1 = cell.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p1.paragraph_format.space_after = Pt(1)
            run = p1.add_run(source)
            run.font.size = Pt(6)
            run.font.color.rgb = self.colors.light_rgb

            # Value
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(1)
            font_size = Pt(11) if label == "PROJECTION" else Pt(18)
            run = p2.add_run(value)
            run.font.size = font_size
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

            # Category label
            p3 = cell.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.paragraph_format.space_before = Pt(0)
            p3.paragraph_format.space_after = Pt(0)
            run = p3.add_run(label)
            run.font.size = Pt(6.5)
            run.font.color.rgb = self.colors.light_rgb

    def _get_stat_value(self, category: str, stat_label: str) -> str:
        """Get a stat value from the appropriate nested stats object."""
        if not self.prospect.stats:
            return "—"

        # Map category names to attribute names
        category_to_attr = {
            "Passing": None,  # PassingStats has flat structure
            "Rushing": "rushing",
            "Receiving": "receiving",
            "Tackles": "tackle",
            "Interceptions": "interception",
        }

        # Normalize the stat label for attribute lookup
        stat_attr = stat_label.lower().replace("%", "_pct")
        if stat_attr == "int":
            stat_attr = "ints"
        elif stat_attr == "tds":
            stat_attr = "td"
        elif stat_attr == "rtg":
            stat_attr = "qb_rtg"

        attr_name = category_to_attr.get(category)

        if attr_name is None:
            # Flat structure (e.g., PassingStats for QB)
            value = getattr(self.prospect.stats, stat_attr, None)
        else:
            # Nested structure (e.g., OffenseSkillPlayerStats.rushing)
            nested_stats = getattr(self.prospect.stats, attr_name, None)
            if nested_stats is None:
                return "—"
            value = getattr(nested_stats, stat_attr, None)

        if value is None:
            return "—"
        return str(value)

    def _gen_stats_bar(self):
        primary_position = get_primary_position(self.prospect.basic_info.position)
        pos_config = POSITION_STATS.get(primary_position, {})
        categories = list(pos_config.keys())

        if categories:
            is_multiple = len(categories) > 1

            if is_multiple:
                total_stats = sum(len(pos_config[cat]) for cat in categories)
                stats_table = self.document.add_table(rows=2, cols=total_stats)
                stats_table.autofit = False

                col_idx = 0
                for cat_idx, category in enumerate(categories):
                    stat_labels = pos_config[category]
                    num_stats = len(stat_labels)
                    bg_color = self.colors.light if cat_idx == 0 else self.colors.medium
                    label_color = "666666" if cat_idx == 0 else self.colors.primary
                    label_rgb = self.color_handler.hex_to_rgb(label_color)

                    # Merge header cells
                    if num_stats > 1:
                        start_cell = stats_table.cell(0, col_idx)
                        end_cell = stats_table.cell(0, col_idx + num_stats - 1)
                        start_cell.merge(end_cell)

                    header_cell = stats_table.cell(0, col_idx)
                    remove_cell_borders(header_cell)
                    set_cell_shading(header_cell, bg_color)
                    set_cell_margins(header_cell, top=80, bottom=40, left=40, right=40)
                    header_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                    p = header_cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(0)
                    run = p.add_run(category.upper())
                    run.font.size = Pt(9)
                    run.font.bold = True
                    run.font.color.rgb = label_rgb

                    for j, stat_label in enumerate(stat_labels):
                        stat_cell = stats_table.cell(1, col_idx + j)
                        stat_cell.width = Inches(6.75 / total_stats)
                        remove_cell_borders(stat_cell)
                        set_cell_shading(stat_cell, bg_color)
                        set_cell_margins(stat_cell, top=80, bottom=40, left=40, right=40)
                        stat_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                        p1 = stat_cell.paragraphs[0]
                        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p1.paragraph_format.space_after = Pt(1)
                        stat_value = self._get_stat_value(category, stat_label)
                        run = p1.add_run(stat_value)
                        run.font.size = Pt(14)
                        run.font.bold = True
                        run.font.color.rgb = self.colors.primary_rgb

                        p2 = stat_cell.add_paragraph()
                        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p2.paragraph_format.space_before = Pt(0)
                        p2.paragraph_format.space_after = Pt(1)
                        run = p2.add_run(stat_label)
                        run.font.size = Pt(7)
                        run.font.color.rgb = label_rgb

                    col_idx += num_stats

            else:
                # Single category (e.g., QB with Passing)
                category = categories[0]
                stat_labels = pos_config[category]

                stats_table = self.document.add_table(rows=1, cols=len(stat_labels))
                stats_table.autofit = False

                for i, stat_label in enumerate(stat_labels):
                    cell = stats_table.cell(0, i)
                    cell.width = Inches(6.75 / len(stat_labels))
                    remove_cell_borders(cell)
                    set_cell_shading(cell, self.colors.light)
                    set_cell_margins(cell, top=100, bottom=40, left=40, right=40)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                    p1 = cell.paragraphs[0]
                    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p1.paragraph_format.space_after = Pt(1)
                    stat_value = self._get_stat_value(category, stat_label)
                    run = p1.add_run(stat_value)
                    run.font.size = Pt(14)
                    run.font.bold = True
                    run.font.color.rgb = self.colors.primary_rgb

                    p2 = cell.add_paragraph()
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p2.paragraph_format.space_before = Pt(0)
                    p2.paragraph_format.space_after = Pt(1)
                    run = p2.add_run(stat_label)
                    run.font.size = Pt(7)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        else:
            # OL - no stats
            p = self.document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run("(Statistics not tracked for Offensive Linemen)")
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    def _gen_skills_and_comps(self):
        skills_table = self.document.add_table(rows=1, cols=2)
        skills_table.autofit = False

        skills_cell = skills_table.cell(0, 0)
        skills_cell.width = Inches(3.8)
        remove_cell_borders(skills_cell)

        skills_header = skills_cell.paragraphs[0]
        skills_header.paragraph_format.space_after = Pt(4)
        run = skills_header.add_run("SKILL RATINGS")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = self.colors.primary_rgb

        # TODO: Skills, comparisons, and recruiting should probably all be distinct methods

        for skill_name, skill_pct in self.prospect.skills.to_dict().items():
            if skill_pct is None:
                continue
            p = skills_cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

            display_name = skill_name.replace("rating", "rtg").replace("targeted", "tgt").replace("_", " ").title()
            run = p.add_run(f"{display_name:<20} ")
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            run = p.add_run(skill_bar(skill_pct))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = self.colors.primary_rgb

            run = p.add_run(f" {skill_pct}%")
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        comp_cell = skills_table.cell(0, 1)
        comp_cell.width = Inches(3.0)
        remove_cell_borders(comp_cell)

        comp_header = comp_cell.paragraphs[0]
        comp_header.paragraph_format.space_after = Pt(4)
        run = comp_header.add_run("PLAYER COMPARISONS")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = self.colors.primary_rgb

        if self.prospect.comparisons is None:
            self.prospect.comparisons = []

        for comp in self.prospect.comparisons:
            p = comp_cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

            run = p.add_run(f"{comp.name} ")
            run.font.size = Pt(9)
            run.font.bold = True

            run = p.add_run(f"({comp.school}) ")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            run = p.add_run(f"{comp.similarity}%")
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = self.colors.primary_rgb

        if self.prospect.ratings:
            p = comp_cell.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run("RECRUITING")
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = self.colors.primary_rgb

            p = comp_cell.add_paragraph()

            run = p.add_run(self.prospect.ratings.get_recruiting_str())
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        self.document.add_paragraph().paragraph_format.space_after = Pt(4)

    def _gen_bio(self):
        if self.prospect.scouting_report.bio:
            header = self.document.add_paragraph()
            header.paragraph_format.space_after = Pt(4)
            run = header.add_run("BACKGROUND")
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = self.colors.primary_rgb

            bio_text = self.prospect.scouting_report.bio.replace("Draft Profile: Bio", "").strip()

            p = self.document.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(bio_text)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # self.document.add_paragraph().paragraph_format.space_after = Pt(4)

    def _gen_strengths_weaknesses(self):
        sw_table = self.document.add_table(rows=1, cols=2)
        sw_table.autofit = False

        str_cell = sw_table.cell(0, 0)
        str_cell.width = Inches(3.375)
        remove_cell_borders(str_cell)

        str_header = str_cell.paragraphs[0]
        str_header.paragraph_format.space_after = Pt(6)
        run = str_header.add_run("STRENGTHS")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1D, 0x6A, 0x4D)

        for strength in self.prospect.scouting_report.strengths:
            p = str_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.15)

            run = p.add_run("+ ")
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1D, 0x6A, 0x4D)

            run = p.add_run(strength)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        weak_cell = sw_table.cell(0, 1)
        weak_cell.width = Inches(3.375)
        remove_cell_borders(weak_cell)

        weak_header = weak_cell.paragraphs[0]
        weak_header.paragraph_format.space_after = Pt(6)
        run = weak_header.add_run("WEAKNESSES")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xA6, 0x5D, 0x21)

        for weakness in self.prospect.scouting_report.weaknesses:
            p = weak_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.15)

            run = p.add_run("– ")
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0xA6, 0x5D, 0x21)

            run = p.add_run(weakness)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    def _gen_scouting_summary(self):
        if self.prospect.scouting_report.summary:
            self.document.add_paragraph().paragraph_format.space_after = Pt(4)

            summary_table = self.document.add_table(rows=1, cols=1)
            summary_table.autofit = False

            cell = summary_table.cell(0, 0)
            cell.width = Inches(6.75)
            set_cell_shading(cell, self.colors.light)
            set_cell_margins(cell, top=100, bottom=70, left=160, right=120)
            add_left_border(cell, self.colors.primary, 24)

            header = cell.paragraphs[0]
            header.paragraph_format.space_after = Pt(4)
            run = header.add_run("SCOUTING SUMMARY")
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = self.colors.primary_rgb

            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(self.prospect.scouting_report.summary.replace("Scouting Report: Summary", ""))
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def generate_complete_document(self):
        self._gen_header_table()

        spacer = self.document.add_paragraph()
        # spacer.paragraph_format.space_after = Pt(2)

        self._gen_rankings_bar()
        self._gen_stats_bar()

        self.document.add_paragraph().paragraph_format.space_after = Pt(10)

        self._gen_skills_and_comps()

        self._gen_bio()
        self.document.add_paragraph().paragraph_format.space_after = Pt(4)

        self._gen_strengths_weaknesses()
        self._gen_scouting_summary()

        full_doc_path = f"{self.output_path}/{self.prospect.basic_info.full_name}.docx"
        self.document.save(full_doc_path)

        if os.path.exists(self.ring_img_path):
            os.remove(self.ring_img_path)

# ─── MAIN BATCH PROCESSOR ────────────────────────────────────────────────────


# def batch_generate_profiles(
#     json_dir: str,
#     photos_dir: str,
#     output_dir: str,
#     colors_file: str,
#     positions: List[str] = None,
#     default_photo: str = None,
# ):
#     """Generate prospect profiles for all positions."""
#     print("Loading school colors...")
#     colors_db = load_school_colors(colors_file)
#     print(f"  Loaded {len(colors_db)} schools")
#
#     os.makedirs(output_dir, exist_ok=True)
#     temp_dir = os.path.join(output_dir, ".temp")
#     os.makedirs(temp_dir, exist_ok=True)
#
#     all_positions = ["QB", "RB", "WR", "TE", "OL", "DL", "EDGE", "LB", "DB"]
#     if positions:
#         all_positions = [p.upper() for p in positions]
#
#     total_generated = 0
#
#     for position in all_positions:
#         json_file = os.path.join(json_dir, f"{position}.json")
#
#         if not os.path.exists(json_file):
#             print(f"Skipping {position}: No JSON file found")
#             continue
#
#         print(f"\nProcessing {position}...")
#
#         with open(json_file, "r") as f:
#             players_data = json.load(f)
#
#         pos_output_dir = os.path.join(output_dir, position)
#         os.makedirs(pos_output_dir, exist_ok=True)
#
#         for rank, (player_name, player_data) in enumerate(players_data.items(), 1):
#             try:
#                 prospect = extract_prospect_data(
#                     player_name, player_data, position, colors_db, photos_dir
#                 )
#
#                 safe_name = (
#                     player_name.replace(" ", "_").replace(".", "").replace("'", "")
#                 )
#                 output_file = os.path.join(
#                     pos_output_dir, f"{rank:02d}_{safe_name}.docx"
#                 )
#
#                 generate_prospect_document(
#                     prospect, output_file, temp_dir, default_photo
#                 )
#
#                 print(
#                     f"  [{rank:3d}] {player_name} ({prospect.school}) - #{prospect.colors['primary']}"
#                 )
#                 total_generated += 1
#
#             except Exception as e:
#                 print(f"  ERROR: {player_name} - {str(e)}")
#                 import traceback
#
#                 traceback.print_exc()
#
#     import shutil
#
#     if os.path.exists(temp_dir):
#         shutil.rmtree(temp_dir)
#
#     print(f"\n{'=' * 50}")
#     print(f"COMPLETE: Generated {total_generated} prospect profiles")
#     print(f"Output directory: {output_dir}")


# if __name__ == "__main__":
#     import argparse
#
#     parser = argparse.ArgumentParser(description="Generate NFL Draft Prospect Profiles")
#     parser.add_argument("--json-dir", required=True)
#     parser.add_argument("--photos-dir", required=True)
#     parser.add_argument("--output-dir", required=True)
#     parser.add_argument("--colors-file", required=True)
#     parser.add_argument("--positions", nargs="+")
#     parser.add_argument("--default-photo")
#
#     args = parser.parse_args()
#
#     batch_generate_profiles(
#         json_dir=args.json_dir,
#         photos_dir=args.photos_dir,
#         output_dir=args.output_dir,
#         colors_file=args.colors_file,
#         positions=args.positions,
#         default_photo=args.default_photo,
#     )
