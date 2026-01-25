#!/usr/bin/env python3
"""
NFL Draft Buzz Prospect Scraper to Word Document

Extracts prospect information from nfldraftbuzz.com and creates a formatted Word document.
Uses Playwright for browser automation to bypass Cloudflare protection.

Usage:
    python nfl_prospect_to_docx.py <url>
    python nfl_prospect_to_docx.py https://www.nfldraftbuzz.com/Player/Dante-Moore-QB-UCLA
    python nfl_prospect_to_docx.py <url> -o custom_output.docx

Requirements:
    pip install playwright python-docx lxml --break-system-packages
    playwright install firefox
"""

import io
import re
import time
from collections import defaultdict
from pathlib import Path
from bs4 import BeautifulSoup, Tag
from playwright.sync_api import sync_playwright, Playwright, Browser, TimeoutError as PlaywrightTimeout, Error as PlaywrightError
from random import uniform
from typing import Optional, List, Dict, Tuple, Union
from urllib.parse import urljoin
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from fbcm.base import POSITION_TO_GROUP_MAP
from fbcm.models import (PassingStats,
                         RushingStats,
                         ReceivingStats,
                         OffenseSkillPlayerStats,
                         DefenseStats,
                         ProspectData,
                         TackleStats,
                         InterceptionStats,
                         Stats,
                         BasicInfo,
                         RatingsAndRankings,
                         Comparison,
                         SkillRatings,
                         PassingSkills,
                         RunningBackSkills,
                         PassCatcherSkills,
                         OffensiveLinemanSkills,
                         DefensiveLinemanSkills,
                         LinebackerSkills,
                         DefensiveBackSkills,
                         ProspectDataSoup, ScoutingReport)


POSITIONS = ["QB", "RB", "WR", "TE", "OL", "DL", "EDGE", "LB", "DB"]


class PageFetcher:
    """Handles fetching web pages using Playwright browser automation."""

    DEFAULT_USER_AGENT = (
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 '
        '(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
    )
    DEFAULT_VIEWPORT = {'width': 1920, 'height': 1080}
    CONTENT_WAIT_TIME = 3000

    MAX_RETRIES = 3

    def __init__(self, playwright: Playwright, headless: bool = False,
                 base_url: str = "https://www.nfldraftbuzz.com"):
        self.base_url = base_url
        self.playwright = playwright
        self.headless = headless
        self.browser = self._launch_browser()

    def _launch_browser(self) -> Browser:
        """Launch a new browser instance."""
        return self.playwright.firefox.launch(headless=self.headless, slow_mo=150)

    def _ensure_browser_connected(self) -> None:
        """Ensure browser is connected, relaunch if necessary."""
        if not self.browser.is_connected():
            print("Browser disconnected, relaunching...")
            self.browser = self._launch_browser()

    def fetch(self, url: str, attempt_image_fetch: bool = False) -> Tuple[str, Optional[bytes], str]:
        """
        Fetch page content using Playwright browser automation.

        Args:
            url: The URL to fetch.
            attempt_image_fetch: Whether to attempt downloading the player image.

        Returns:
            Tuple of (page_text, image_bytes, image_type).
        """
        last_error = None

        for attempt in range(self.MAX_RETRIES):
            try:
                return self._fetch_with_page(url, attempt_image_fetch)
            except PlaywrightError as e:
                last_error = e
                error_msg = str(e).lower()
                if "target closed" in error_msg or "browser has been closed" in error_msg:
                    print(f"Browser/target closed (attempt {attempt + 1}/{self.MAX_RETRIES}), relaunching...")
                    try:
                        self.browser.close()
                    except Exception:
                        pass  # Browser may already be closed
                    self.browser = self._launch_browser()
                    time.sleep(1)  # Brief pause before retry
                else:
                    raise  # Re-raise if it's a different Playwright error

        raise last_error  # All retries exhausted

    def fetch_soup(self, url) -> BeautifulSoup:
        self._ensure_browser_connected()
        page = self.browser.new_page()
        try:
            print(f"Navigating to: {url}")
            page.goto(url=url)
            return BeautifulSoup(page.content(), "lxml")
        finally:
            page.close()

    def _fetch_with_page(self, url: str, attempt_image_fetch: bool) -> Tuple[str, Optional[bytes], str]:
        """Internal method to fetch a page. May raise PlaywrightError."""
        self._ensure_browser_connected()
        print("Opening new page...")

        page = self.browser.new_page()
        try:
            print(f"Navigating to: {url}")
            try:
                page.goto(url)
            except PlaywrightTimeout:
                print("Page load timeout, continuing with partial content...")

            text_content = page.evaluate('() => document.body.innerText')
            if attempt_image_fetch:
                image_data, image_type = self._find_and_download_image(page, url)
            else:
                image_data = None
                image_type = None
            # TODO: Returning both text_content and page.content is a temporary kludge
            return text_content, image_data, image_type
        finally:
            page.close()

    def _find_and_download_image(self, page, base_url: str) -> Tuple[Optional[bytes], str]:
        """Find and download the player image from the page."""
        image_url = self._find_image_url(page)

        if not image_url:
            image_url = self._find_any_large_image(page)

        if image_url:
            return self._download_image(page, image_url, base_url)

        return None, "jpeg"

    def _find_image_url(self, page) -> Optional[str]:
        """Try to find image URL using predefined selectors."""
        img = page.locator("figure.player-info__photo img")
        src = img.get_attribute("src")
        return self._make_absolute_url(url=src, base_url=self.base_url)

    def _find_any_large_image(self, page) -> Optional[str]:
        """Fallback: try to find any large player image."""
        try:
            images = page.query_selector_all('img')
            for img in images:
                src = img.get_attribute('src')
                if src and not self._should_skip_image(src):
                    if 'nfldraftbuzz' in src or 'imagn' in src.lower() or 'player' in src.lower():
                        return src
        except Exception:
            pass
        return None

    def _should_skip_image(self, src: str) -> bool:
        """Check if an image URL should be skipped."""
        src_lower = src.lower()
        return any(pattern in src_lower for pattern in self.SKIP_IMAGE_PATTERNS)

    def _download_image(self, page, image_url: str, base_url: str) -> Tuple[Optional[bytes], str]:
        """Download image from URL."""
        print(f"Found player image: {image_url[:80]}...")
        try:
            image_url = self._make_absolute_url(image_url, base_url)
            response = page.request.get(image_url)
            if response.ok:
                image_data = response.body()
                image_type = self._get_image_type(response.headers.get('content-type', ''))
                print(f"Downloaded image: {len(image_data)} bytes ({image_type})")
                return image_data, image_type
        except Exception as e:
            print(f"Failed to download image: {e}")
        return None, "jpeg"

    @staticmethod
    def _make_absolute_url(url: str, base_url: str = None) -> str:
        """Convert relative URL to absolute."""
        if url.startswith('//'):
            return 'https:' + url
        elif url.startswith('/'):
            return urljoin(base_url, url)
        return url

    @staticmethod
    def _get_image_type(content_type: str) -> str:
        """Determine image type from content-type header."""
        if 'png' in content_type:
            return 'png'
        elif 'gif' in content_type:
            return 'gif'
        elif 'webp' in content_type:
            return 'webp'
        return 'jpeg'


class ProspectParserSoup:
    """
    Parses nfldraftbuzz.com prospect profiles using BeautifulSoup
    """
    def __init__(self, soup: BeautifulSoup, position: str):
        self.soup = soup
        self.position = position

    ##### Utility Methods #####
    def _get_tag_with_title_containing(self, tag, search_str) -> Tag:
        return tag.find("span", title=lambda t: t and search_str in t)

    def _get_tag_with_text(self, search_space, tag_name, text):
        # Note that text should be lower case since we use lower()
        return search_space.find(tag_name, string=lambda s: text in s.lower())

    def _get_text_following_label(self, label_tag, expected_sibling_name: str = "span") -> str:
        return label_tag.find_next_sibling(expected_sibling_name).get_text(strip=True)

    ##### Basic Info Related #####
    def parse_basic_info(self) -> BasicInfo:
        basic_info_dict = {}

        first_name, last_name = self._parse_name()

        info_details_div = self.soup.find("div", class_="player-info-details")
        basic_info_dict.update(self._parse_player_info_details_div(div=info_details_div))

        basic_info_table_tag = self.soup.find("table", class_="basicInfoTable")
        basic_info_dict.update(self._parse_basic_info_table(basic_info_table_tag))

        basic_info_dict["class_"] = basic_info_dict.pop("class")
        basic_info_dict["hometown"] = basic_info_dict.pop("home town")

        return BasicInfo(first_name=first_name,
                         last_name=last_name,
                         full_name=f"{first_name} {last_name}",
                         **basic_info_dict)

    def parse_ratings(self, table: Tag) -> RatingsAndRankings:
        self._perform_rating_checks(table=table)

        table_rows = table.find_all("tr")
        overall = self._extract_ovr_rtg(row=table_rows[0])
        opposition = self._extract_opposition_rtg(row=table_rows[2])

        proj_rank_row = self._get_projection_ranks_row(rows=table_rows)
        proj_ranks = self._extract_proj_and_rankings(row=proj_rank_row)

        game_snap_count_row = table_rows[8]

        outlet_ratings = self._extract_outlet_ratings(table=table)

        rate_ranks = RatingsAndRankings(
            overall_rating=overall,
            opposition_rating=opposition,
            **proj_ranks,
            **outlet_ratings
        )

        return rate_ranks

    def parse_skills(self, table: Tag) -> SkillRatings:
        rows = table.find_all("tr")[4:]
        skill_rtgs_rows = self._gather_skill_rtg_rows(rows=rows)
        skill_ratings_dict = self._extract_skill_ratings(rows=skill_rtgs_rows)
        return self._construct_skill_ratings_obj(ratings=skill_ratings_dict)

    def parse_comparisons(self, table: Tag) -> List[Comparison]:
        comparisons = []
        comp_rows = table.find("tbody").find_all("tr")

        for row in comp_rows:
            text_parts = row.get_text().split()
            comp_name = f"{text_parts[0]} {text_parts[1]}"
            comp_school = text_parts[3]
            comp_score = int(text_parts[-1].replace("%", ""))

            comparisons.append(Comparison(name=comp_name,
                                          school=comp_school,
                                          similarity=comp_score))

        return comparisons

    def parse_scouting_report(self) -> ScoutingReport:
        intro_div = self.soup.find("div", class_="playerDescIntro")
        strengths_div = self.soup.find("div", class_="playerDescPro")
        weaknesses_div, summary_div = self.soup.find_all("div", class_="playerDescNeg")

        strengths = [line for line in strengths_div.get_text().splitlines()
                     if line and "scouting report" not in line.lower()]
        weaknesses = [line for line in weaknesses_div.get_text().splitlines()
                      if line and "scouting report" not in line.lower()]

        return ScoutingReport(bio=intro_div.get_text(strip=True),
                              strengths=strengths,
                              weaknesses=weaknesses,
                              summary=summary_div.get_text(strip=True))

    def extract_image_url(self) -> str:
        figure_tag = self.soup.find("figure",
                                    class_="player-info__photo")
        image_path = figure_tag.find("img")["src"]
        return f"https://www.nfldraftbuzz.com{image_path}"

    def parse_stats(self, soup: BeautifulSoup) -> Stats:

        stats = None
        table_div = None
        match self.position:
            case "QB":
                table_div = soup.find(id="QBstats")
            case "RB" | "WR" | "TE":
                table_div = soup.find(id="RB-Rush-stats")
            case "OL":
                pass
            case "DL" | "EDGE" | "LB" | "DB":
                table_div = soup.find(id="DBLBDL-stats")
            case _:
                print(f"Could not match position {self.position} to any known group.")

        if table_div is not None:
            stats = self._extract_stats_object(div=table_div)[0]

        return stats

    def parse(self):
        basic_info = self.parse_basic_info()
        rtgs_table, comps_table = self._extract_ratings_comps_tables()

        ratings = self.parse_ratings(table=rtgs_table)
        skills = self.parse_skills(table=rtgs_table)
        comparisons = self.parse_comparisons(table=comps_table)
        scouting_report = self.parse_scouting_report()

        return ProspectDataSoup(
            basic_info=basic_info,
            ratings=ratings,
            skills=skills,
            comparisons=comparisons,
            scouting_report=scouting_report,
            stats=None
        )

    ##### Basic Info #####
    def _parse_name(self) -> Tuple[str, str]:
        first_name = self.soup.find("span", class_="player-info__first-name").get_text(strip=True)
        last_name = self.soup.find("span", class_="player-info__last-name").get_text(strip=True)

        return first_name, last_name

    def _parse_position(self, value: str) -> str:
        position_group_str = ""
        if "/" in value:
            p1, p2 = value.split("/")
            p1_group = POSITION_TO_GROUP_MAP.get(p1)
            p2_group = POSITION_TO_GROUP_MAP.get(p2)

            # Neither correspond to a known group
            if not (p1_group or p2_group):
                raise ValueError(f"Could not find a valid position group for position: {value}")

            # Both correspond to a known group
            if p1_group and p2_group:
                position_group_str = f"{p1_group}/{p2_group}"
            # Only one of the two values corresponds to a known group
            # p1_group is the "real" one
            elif p1_group:
                position_group_str = p1_group
            # If only one value corresponds to a real group, and it's not p1
            # it must be p2
            elif p2:
                position_group_str = p2_group

        else:
            position_group_str = POSITION_TO_GROUP_MAP[value.upper()]

        return position_group_str

    def _parse_player_info_details_div(self, div: Tag) -> Dict:
        # This div contains the values for:
        # height, weight, college, position, player_class, hometown
        basic_info_dict = {}

        for attr_div in div.find_all("div", class_="player-info-details__item"):
            field_tag = attr_div.find("h6", class_="player-info-details__title")
            value_tag = attr_div.find("div", class_="player-info-details__value")

            field = field_tag.get_text(strip=True).lower()
            value = value_tag.get_text(strip=True).lower()

            if field == "position":
                value = self._parse_position(value=value)
            basic_info_dict[field] = value

        return basic_info_dict

    def _parse_basic_info_table(self, tag: Tag) -> Dict:
        # Includes jersery #, sub_position, last_updated, forty_time
        jersey_num = tag.find(text=re.compile(r"#\d+")).get_text(strip=True)

        sub_position_label = self._get_tag_with_title_containing(tag, "Sub-Position")
        sub_position_value = self._get_text_following_label(sub_position_label)

        last_updated_label = self._get_tag_with_title_containing(tag, "Last Updated")
        last_updated_value = self._get_text_following_label(last_updated_label)

        draft_yr_label = self._get_tag_with_title_containing(tag, "Draft Year")
        draft_yr_value = self._get_text_following_label(draft_yr_label)

        forty_label = self._get_tag_with_title_containing(tag, "40 yard dash time")
        forty_value = self._get_text_following_label(forty_label)

        return {
            "jersey": jersey_num,
            "play_style": sub_position_value,
            "last_updated": last_updated_value,
            "draft_year": draft_yr_value,
            "forty": forty_value
        }

    ##### Statistical Related #####
    def _transform_passing_stats(self, season_stats):
        season_stats["cmp_pct"] = season_stats.pop("cmp%")
        season_stats["ints"] = season_stats.pop("int")
        season_stats["qb_rtg"] = season_stats.pop("pro rat")
        season_stats.pop("rat")
        season_stats.pop("avg")

        season_stats["year"] = season_stats.pop("year").split()[0]

        from pprint import pprint
        for fld in ["cmp", "att", "yds", "td", "ints", "sack", "year"]:
            try:
                season_stats[fld] = int(season_stats[fld] or 0)
            except ValueError as e:
                print(f"Invalid value for field {fld}: {season_stats[fld]}")
                print("Full season_stats_dict")
                pprint(season_stats, indent=4)
                raise e

        for fld in ["cmp_pct", "qb_rtg"]:
            try:
                season_stats[fld] = float(season_stats[fld] or 0.0)
            except ValueError as e:
                print(f"Invalid value for field {fld}: {season_stats[fld]}")
                print("Full season_stats_dict")
                pprint(season_stats, indent=4)
                raise e

        return season_stats

    def _transform_stats(self, season_stats):
        match self.position:
            case "QB":
                return self._transform_passing_stats(season_stats=season_stats)
            case "RB":
                pass
            case "WR":
                pass
            case "TE":
                pass
            case "OL":
                pass
            case "DL":
                pass
            case "EDGE":
                pass
            case "LB":
                pass
            case "DB":
                pass
        return season_stats

    def _extract_stats_object(self, div):
        stats_table = div.find("table")
        header_values = [th.get_text(strip=True).lower()
                         for th in stats_table.thead.find_all("th", class_="player-season-avg__stat")
                         if th.get_text(strip=True)]
        seasons = []

        for row in stats_table.tbody.find_all("tr"):
            cells = [td.get_text(strip=True) for td in row.find_all("td")]

            if self.position == "QB":
                season_stats = dict(zip(header_values, cells))
                season_stats = self._transform_stats(season_stats=season_stats)
                stats_obj = PassingStats(**season_stats)
            elif self.position in ["RB"]:
                season_stats = {
                    "year": cells[0],
                    "rushing": {
                        "att": int(cells[1]),
                        "yds": int(cells[2]),
                        "avg": float(cells[3]),
                        "td": int(cells[4])
                    },
                    "receiving": {
                        "rec": int(cells[5]),
                        "yds": int(cells[6]),
                        "avg": float(cells[7]),
                        "td": int(cells[8])
                    }
                }
                rushing_stats = RushingStats(year=season_stats["year"],
                                             **season_stats["rushing"])
                receiving_stats = ReceivingStats(year=season_stats["year"],
                                                 **season_stats["receiving"])
                stats_obj = OffenseSkillPlayerStats(year=season_stats["year"],
                                                    rushing=rushing_stats,
                                                    receiving=receiving_stats)
            elif self.position in ["WR", "TE"]:
                season_stats = {
                    "year": cells[0],
                    "receiving": {
                        "rec": int(cells[1]),
                        "yds": int(cells[2]),
                        "avg": float(cells[3]),
                        "td": int(cells[4])
                    },
                    "rushing": {
                        "att": int(cells[5]),
                        "yds": int(cells[6]),
                        "avg": float(cells[7]),
                        "td": int(cells[8])
                    }
                }
                rushing_stats = RushingStats(year=season_stats["year"],
                                             **season_stats["rushing"])
                receiving_stats = ReceivingStats(year=season_stats["year"],
                                                 **season_stats["receiving"])
                stats_obj = OffenseSkillPlayerStats(year=season_stats["year"],
                                                    rushing=rushing_stats,
                                                    receiving=receiving_stats)
            elif self.position == "OL":
                stats_obj = None
            elif self.position in ["DL", "EDGE", "LB", "DB"]:
                season_stats = {
                    "year": int(cells[0].split()[0]),
                    "tackle": {
                        "total": int(cells[1]),
                        "solo": int(cells[2]),
                        "ff": int(cells[3]),
                        "sacks": float(cells[4])
                    },
                    "interception": {
                        "ints": int(cells[5]),
                        "yds": int(cells[6]),
                        "td": int(cells[7]),
                        "pds": int(cells[8])
                    }
                }
                tackle_stats = TackleStats(year=season_stats["year"],
                                           **season_stats["tackle"])
                interception_stats = InterceptionStats(year=season_stats["year"],
                                                       **season_stats["interception"])
                stats_obj = DefenseStats(year=season_stats["year"],
                                         tackle=tackle_stats,
                                         interception=interception_stats)

            else:
                raise ValueError(f"Could not match position {self.position} to "
                                 f"a position with a defined stats mapping.")

            seasons.append(stats_obj)

        seasons.sort(key=lambda stats: stats.year, reverse=True)

        return seasons

    ##### Ratings and Grades #####
    def _perform_rating_checks(self, table: Tag):
        ovr_rtg_label = table.find("th")
        if "overall rating" not in ovr_rtg_label.get_text().lower():
            raise ValueError(f"Unexpected label in first <th> element: {ovr_rtg_label.get_text}")

    def _extract_ovr_rtg(self, row: Tag) -> float:
        ovr_rtg = float(row.find("span")
                        .get_text(strip=True)
                        .replace(" / 100", ""))
        return ovr_rtg

    def _extract_opposition_rtg(self, row: Tag) -> int:
        meter_div = row.find("div", class_="meter")
        rtg_as_str = meter_div["title"].split(":")[-1].strip().replace("%", "")
        return int(rtg_as_str)

    def _extract_skill_ratings(self, rows: List[Tag]) -> Dict:
        skills = {}
        for row in rows:
            skill_name, rating = (row.get_text(strip=True).lower()
                                  .replace(" ", "_").replace("%", "").split(":"))


            if "/" in rating:
                rating = rating.split("/")[0]
            rating = float(rating.replace("_", ""))

            skills[skill_name.replace("/", "_")] = int(rating)

        return skills

    def _extract_proj_and_rankings(self, row) -> Dict:
        projection_label = self._get_tag_with_text(search_space=row,
                                                   tag_name="span",
                                                   text="draft projection")
        projection = self._get_text_following_label(label_tag=projection_label)

        ovr_rank_label = self._get_tag_with_text(search_space=row,
                                           tag_name="span",
                                           text="overall rank")
        ovr_rank = self._get_text_following_label(label_tag=ovr_rank_label)

        pos_rank_label = self._get_tag_with_text(search_space=row,
                                                 tag_name="span",
                                                 text="position rank")
        pos_rank = self._get_text_following_label(label_tag=pos_rank_label)

        return {
            "draft_projection": projection,
            "overall_rank": ovr_rank,
            "position_rank": pos_rank
        }

    def _get_projection_ranks_row(self, rows: List[Tag]) -> Optional[Tag]:
        for row in rows:
            if "draft projection" in row.get_text().lower():
                return row
        return None

    def _gather_skill_rtg_rows(self, rows: List[Tag], sentinel_val: str = "draft projection") -> List[Tag]:
        skill_rows = []
        for row in rows:
            if sentinel_val in row.get_text().lower():
                break
            skill_rows.append(row)

        return skill_rows

    def _construct_skill_ratings_obj(self, ratings: Dict) -> SkillRatings:
        skills = None
        match self.position:
            case "QB":
                skills = PassingSkills(**ratings)
            case "RB":
                skills = RunningBackSkills(**ratings)
            case "WR" | "TE":
                skills = PassCatcherSkills(**ratings)
            case "OL":
                skills = OffensiveLinemanSkills(**ratings)
            case "DL" | "EDGE":
                skills = DefensiveLinemanSkills(**ratings)
            case "LB":
                skills = LinebackerSkills(**ratings)
            case "DB":
                skills = DefensiveBackSkills(**ratings)
            case _:
                raise ValueError(f"Could not find skill ratings for position: {self.position}")
        return skills

    ##### Outlet ratings ####
    def _extract_outlet_ratings(self, table: Tag) -> Dict[str, Optional[float]]:
        return {
            "espn": self._extract_espn(table=table),
            "rivals": self._extract_rivals(table=table),
            "rtg_247": self._extract_247(table=table)
        }

    def _extract_rivals(self, table: Tag) -> Optional[float]:
        rivals_row = self._get_tag_with_text(search_space=table,
                                             tag_name="span",
                                             text="rivals")
        if rivals_row:
            rivals_rtg = float(rivals_row.get_text(strip=True).split(":")[-1].split()[0])
        else:
            rivals_rtg = None

        return rivals_rtg

    def _extract_247(self, table: Tag) -> Optional[float]:
        rtg = None
        sports_247_rtg_row = self._get_tag_with_text(search_space=table,
                                                     tag_name="span",
                                                     text="247")
        if sports_247_rtg_row:
            rtg = float(sports_247_rtg_row.get_text(strip=True).split()[-1].split("/")[0])

        return rtg

    def _extract_espn(self, table: Tag) -> Optional[float]:
        rtg = None
        espn_rtg_row = self._get_tag_with_text(search_space=table,
                                                 tag_name="span",
                                                 text="espn")
        if espn_rtg_row:
            rtg = float(espn_rtg_row.get_text(strip=True).split()[-1].split("/")[0])

        return rtg

    def _extract_ratings_comps_tables(self):
        ratings_and_rankings = [table for table
                                in self.soup.find_all("table", class_="starRatingTable")
                                if not table.find("th", string=lambda s: "measurables" in s.lower())]

        ratings = ratings_and_rankings[0]
        comparisons = ratings_and_rankings[1]
        return ratings, comparisons


class ProspectParser:
    """Parses page text content and extracts prospect data."""

    POSITIONS = {'QB', 'RB', 'WR', 'TE', 'OL', 'DL', 'EDGE', 'LB', 'DB', 'CB', 'S', 'PK', 'P'}

    SKILL_PATTERNS = [
        ('Release Speed', r'RELEASE SPEED:\s*\n?\s*(\d+)%'),
        ('Short Passing', r'SHORT PASSING:\s*\n?\s*(\d+)%'),
        ('Medium Passing', r'MEDIUM PASSING:\s*\n?\s*(\d+)%'),
        ('Long Passing', r'LONG PASSING:\s*\n?\s*(\d+)%'),
        ('Rush/Scramble', r'RUSH/SCRAMBLE:\s*\n?\s*(\d+)%'),
    ]

    def __init__(self):
        self.soup = None
        self.position = None

    def parse(self, text: str, image_data: Optional[bytes] = None,
              image_type: str = "jpeg", html_str: str = None) -> ProspectData:
        """Parse page text and extract all prospect data."""
        data = ProspectData()
        data.image_data = image_data
        data.image_type = image_type



        self._parse_name(text, data)
        self._parse_basic_info(text, data)
        self._parse_ratings(text, data)
        # self._parse_stats(text, data)
        self._parse_skill_ratings(text, data)
        self._parse_recruiting_grades(text, data)
        self._parse_comparisons(text, data)
        self._parse_scouting_content(text, data)
        self._parse_consensus_rankings(text, data)

        return data

    def _parse_name(self, text: str, data: ProspectData) -> None:
        """Extract player name from page text."""
        pos_pattern = '|'.join(self.POSITIONS)
        title_match = re.search(
            rf'([A-Z]+)\s+([A-Z]+)\s+({pos_pattern})\s+[A-Z]+\s*\|\s*NFL DRAFT',
            text
        )
        if title_match:
            data.name = f"{title_match.group(1).title()} {title_match.group(2).title()}"
            return

        name_match = re.search(r'SIMULATOR\s*\n?\s*([A-Z]+)\s*\n\s*([A-Z]+)\s*\n?\s*HEIGHT', text)
        if name_match:
            data.name = f"{name_match.group(1).title()} {name_match.group(2).title()}"

    def _parse_basic_info(self, text: str, data: ProspectData) -> None:
        # TODO: All but college_games and college_snaps are completed.
        """Extract basic player information."""
        patterns = {
            'height': (r'HEIGHT\s*\n?\s*(\d+-\d+)', 'height'),
            'weight': (r'WEIGHT\s*\n?\s*(\d+)', 'weight'),
            'school': (r'COLLEGE\s*\n?\s*([A-Za-z\s]+?)\s*\n', 'school'),
            'position': (rf'POSITION\s*\n?\s*({"|".join(self.POSITIONS)})', 'position'),
            'player_class': (r'CLASS\s*\n?\s*(Freshman|Sophomore|Junior|Senior|RS\s*\w+)', 'player_class'),
            'hometown': (r'HOME\s*TOWN\s*\n?\s*([A-Za-z\s,]+?)(?:\n|$)', 'hometown'),
            'jersey': (r'JERSEY:\s*#(\d+)', 'jersey'),
            'play_style': (r'PLAY STYLE:\s*([A-Z\s]+?)(?:\n|LAST)', 'play_style'),
            'last_updated': (r'LAST UPDATED:\s*(\d+/\d+/\d+)', 'last_updated'),
            'draft_year': (r'DRAFT YEAR:\s*(\d+)', 'draft_year'),
            'college_games': (r'COLLEGE GAMES:\s*(\d+)', 'college_games'),
            'college_snaps': (r'COLLEGE SNAPS:\s*(\d+)', 'college_snaps'),
        }

        for _, (pattern, attr) in patterns.items():
            match = re.search(pattern, text, re.IGNORECASE if attr == 'player_class' else 0)
            if match:
                value = match.group(1).strip() if attr in ('school', 'hometown', 'play_style') else match.group(1)
                setattr(data, attr, value)

        print("ARMADILLO")
        print(data)

        self.position = data.position

        # Age and DOB combined pattern
        age_match = re.search(r'AGE:\s*([\d.]+)\s*DOB:\s*(\d+/\d+/\d+)', text)
        if age_match:
            data.age = age_match.group(1)
            data.dob = age_match.group(2)

        # Height/Weight with percentiles
        hw_match = re.search(r'HEIGHT:\s*(\d+-\d+)\s*\((\d+)%\*?\)\s*WEIGHT:\s*(\d+)\s*\((\d+)%', text)
        if hw_match:
            data.height = f"{hw_match.group(1)} ({hw_match.group(2)}%)"
            data.weight = f"{hw_match.group(3)} ({hw_match.group(4)}%)"

        # Forty time with percentile
        forty_match = re.search(r'FORTY TIME:\s*([\d.]+)\s*SECONDS\s*\((\d+)%', text)
        if forty_match:
            data.forty = f"{forty_match.group(1)} ({forty_match.group(2)}%)"
        elif not data.forty:
            forty_simple = re.search(r'(\d+\.\d+)\s*\n?\s*FORTY\s*YD\s*TIME', text)
            if forty_simple:
                data.forty = forty_simple.group(1)

    def _parse_ratings(self, text: str, data: ProspectData) -> None:
        """Extract player ratings."""
        patterns = {
            'overall_rating': r'(\d+\.\d+)\s*/100\s*\n?\s*PLAYER\s*RATING',
            'position_rank': r'(\d+)\s*\n?\s*POSITION\s*RANK',
            'overall_rank': r'OVERALL RANK:\s*#(\d+)',
            'draft_projection': r'DRAFT PROJECTION:\s*([^\n]+)',
            'defense_rating': r'DEFENSE RATING:\s*\n?\s*(\d+)%',
        }

        for attr, pattern in patterns.items():
            match = re.search(pattern, text)
            if match:
                value = match.group(1).strip() if attr == 'draft_projection' else match.group(1)
                setattr(data, attr, value)

    def _parse_stats(self, text: str, data: ProspectData) -> None:
        """Extract player statistics (primarily for QBs)."""
        patterns = {
            'qb_rating': r'QB RATING\s*\n?\s*([\d.]+)',
            'yards': r'YDS\s*\n?\s*(\d+)',
            'comp_pct': r'COMP %\s*\n?\s*([\d.]+)',
            'tds': r'TDS\s*\n?\s*(\d+)',
            'ints': r'INTS\s*\n?\s*(\d+)',
            'rush_avg': r'RUSH AVG\s*\n?\s*([\d.]+)',
        }

        for attr, pattern in patterns.items():
            match = re.search(pattern, text)
            if match:
                setattr(data, attr, match.group(1))

    def _parse_skill_ratings(self, text: str, data: ProspectData) -> None:
        """Extract skill ratings (percentiles)."""
        for skill_name, pattern in self.SKILL_PATTERNS:
            match = re.search(pattern, text)
            if match:
                data.skill_ratings[skill_name] = match.group(1)

    def _parse_recruiting_grades(self, text: str, data: ProspectData) -> None:
        """Extract recruiting grades from various services."""
        espn_match = re.search(r'ESPN RATING:\s*([\d/]+)', text)
        if espn_match:
            data.espn_rating = espn_match.group(1)

        rating247_match = re.search(r'247 RATING:\s*([\d/]+)', text)
        if rating247_match:
            data.rating_247 = rating247_match.group(1)

        rivals_match = re.search(r'RIVALS RATING:\s*([\d.]+\s*\([^)]+\))', text)
        if rivals_match:
            data.rivals_rating = rivals_match.group(1)

    def _parse_comparisons(self, text: str, data: ProspectData) -> None:
        """Extract player comparisons."""
        comp_pattern = r'([A-Z][a-z]+)\s+([A-Z]+)\s*-\s*([A-Z\s]+?)\s*\n?\s*(\d+)%'
        for match in re.finditer(comp_pattern, text):
            name = f"{match.group(1)} {match.group(2)}"
            school = match.group(3).strip()
            similarity = match.group(4)
            data.comparisons.append((name, school, similarity))

    def _parse_scouting_content(self, text: str, data: ProspectData) -> None:
        """Extract scouting report content."""
        # Bio
        bio_match = re.search(r'DRAFT PROFILE: BIO\s*\n(.+?)(?=SCOUTING REPORT: STRENGTHS)', text, re.DOTALL)
        if bio_match:
            data.bio = bio_match.group(1).strip()

        # Strengths
        strengths_match = re.search(
            r'SCOUTING REPORT: STRENGTHS\s*\n(.+?)(?=SCOUTING REPORT: WEAKNESSES)', text, re.DOTALL
        )
        if strengths_match:
            data.strengths = self._split_scouting_points(strengths_match.group(1))

        # Weaknesses
        weaknesses_match = re.search(
            r'SCOUTING REPORT: WEAKNESSES\s*\n(.+?)(?=SCOUTING REPORT: SUMMARY)', text, re.DOTALL
        )
        if weaknesses_match:
            data.weaknesses = self._split_scouting_points(weaknesses_match.group(1))

        # Summary
        summary_match = re.search(r'SCOUTING REPORT: SUMMARY\s*\n(.+?)(?=NEXT:|HOW OTHER)', text, re.DOTALL)
        if summary_match:
            data.summary = summary_match.group(1).strip()

    @staticmethod
    def _split_scouting_points(text: str) -> List[str]:
        """Split scouting text into individual points."""
        points = [s.strip() for s in re.split(r'\.\s+(?=[A-Z])', text.strip()) if s.strip()]
        return [s if s.endswith('.') else s + '.' for s in points]

    def _parse_consensus_rankings(self, text: str, data: ProspectData) -> None:
        """Extract consensus rankings from other scouts."""
        avg_overall_match = re.search(r'ALL SCOUTS AVERAGE\s*OVERALL RANK\s*\n?\s*([\d.]+)', text)
        if avg_overall_match:
            data.avg_overall_rank = avg_overall_match.group(1)

        avg_pos_match = re.search(r'ALL SCOUTS AVERAGE\s*POSITION RANK\s*\n?\s*([\d.]+)', text)
        if avg_pos_match:
            data.avg_position_rank = avg_pos_match.group(1)

    def extract_name_from_url(self, url: str) -> Optional[str]:
        """Extract player name from URL as fallback."""
        url_match = re.search(r'/Player/([^/]+)', url)
        if url_match:
            name_parts = url_match.group(1).replace('-', ' ').split()
            name_parts = [p for p in name_parts if p.upper() not in self.POSITIONS]
            if len(name_parts) >= 2:
                return ' '.join(name_parts[:2])
        return None


class WordDocumentGenerator:
    """Generates Word documents from prospect data."""

    HEADER_COLOR = "003366"
    STATS_HEADER_COLOR = "006633"
    SKILLS_HEADER_COLOR = "663399"
    COMPARISONS_HEADER_COLOR = "CC6600"

    def __init__(self):
        self.doc = None

    def generate(self, data: ProspectData, output_path: str) -> None:
        """Create a formatted Word document from prospect data."""
        self.doc = Document()
        self._setup_styles()

        self._add_title(data)
        self._add_player_image(data)
        self._add_player_information(data)
        self._add_ratings(data)
        self._add_statistics(data)
        self._add_skill_ratings(data)
        self._add_recruiting_grades(data)
        self._add_comparisons(data)
        self._add_bio(data)
        self._add_strengths(data)
        self._add_weaknesses(data)
        self._add_summary(data)
        self._add_consensus_rankings(data)
        self._add_footer(data)

        self.doc.save(output_path)
        print(f"Document saved to: {output_path}")

    def _setup_styles(self) -> None:
        """Set up document styles."""
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

    def _add_title(self, data: ProspectData) -> None:
        """Add title and subtitle to document."""
        title = self.doc.add_heading(level=0)
        title_run = title.add_run(f"{data.name.upper()}")
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = self.doc.add_paragraph()
        subtitle_run = subtitle.add_run(f"{data.position} | {data.school}")
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph("─" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_player_image(self, data: ProspectData) -> None:
        """Add player image if available."""
        if not data.image_data:
            return

        try:
            image_para = self.doc.add_paragraph()
            image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = image_para.add_run()
            image_stream = io.BytesIO(data.image_data)
            run.add_picture(image_stream, width=Inches(4))

            caption = self.doc.add_paragraph()
            caption_run = caption.add_run(f"{data.name} - {data.school}")
            caption_run.font.size = Pt(10)
            caption_run.font.italic = True
            caption_run.font.color.rgb = RGBColor(102, 102, 102)
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

            self.doc.add_paragraph()
        except Exception as e:
            print(f"Warning: Could not add image to document: {e}")

    def _add_player_information(self, data: ProspectData) -> None:
        """Add player information table."""
        self.doc.add_heading("PLAYER INFORMATION", level=1)

        info_items = [
            ("Jersey", f"#{data.jersey}" if data.jersey else "N/A"),
            ("Play Style", data.play_style or "N/A"),
            ("Height", data.height or "N/A"),
            ("Weight", data.weight or "N/A"),
            ("40-Yard Dash", data.forty or "N/A"),
            ("Age", data.age or "N/A"),
            ("DOB", data.dob or "N/A"),
            ("Hometown", data.hometown or "N/A"),
            ("Class", data.player_class or "N/A"),
            ("Draft Year", data.draft_year or "N/A"),
            ("College Games", data.college_games or "N/A"),
            ("College Snaps", data.college_snaps or "N/A"),
        ]

        info_table = self.doc.add_table(rows=0, cols=4)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i in range(0, len(info_items), 2):
            row = info_table.add_row()
            row.cells[0].text = info_items[i][0] + ":"
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = info_items[i][1]
            if i + 1 < len(info_items):
                row.cells[2].text = info_items[i + 1][0] + ":"
                row.cells[2].paragraphs[0].runs[0].bold = True
                row.cells[3].text = info_items[i + 1][1]

        self.doc.add_paragraph()

    def _add_ratings(self, data: ProspectData) -> None:
        """Add ratings and rankings table."""
        self.doc.add_heading("RATINGS & RANKINGS", level=1)

        ratings_table = self.doc.add_table(rows=2, cols=4)
        ratings_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["Overall Rating", "Position Rank", "Overall Rank", "Draft Projection"]
        values = [
            f"{data.overall_rating}/100" if data.overall_rating else "N/A",
            f"#{data.position_rank} ({data.position})" if data.position_rank else "N/A",
            f"#{data.overall_rank}" if data.overall_rank else "N/A",
            data.draft_projection or "N/A"
        ]

        for i, header in enumerate(headers):
            cell = ratings_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_shading(cell, self.HEADER_COLOR)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for i, value in enumerate(values):
            cell = ratings_table.rows[1].cells[i]
            cell.text = value
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_shading(cell, "E6F2FF")

        self.doc.add_paragraph()

    def _add_statistics(self, data: ProspectData) -> None:
        """Add statistics table (for QBs)."""
        if data.position != "QB" or not data.stats:
            return

        self.doc.add_heading("STATISTICS", level=1)

        stats_table = self.doc.add_table(rows=2, cols=6)
        stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        stat_headers = ["QB Rating", "Yards", "Comp %", "TDs", "INTs", "Rush Avg"]
        stat_values = [
            data.qb_rating or "N/A",
            data.yards or "N/A",
            f"{data.comp_pct}%" if data.comp_pct else "N/A",
            data.tds or "N/A",
            data.ints or "N/A",
            data.rush_avg or "N/A"
        ]

        for i, header in enumerate(stat_headers):
            cell = stats_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_shading(cell, self.STATS_HEADER_COLOR)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for i, value in enumerate(stat_values):
            cell = stats_table.rows[1].cells[i]
            cell.text = value
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_shading(cell, "E6FFE6")

        self.doc.add_paragraph()

    def _add_skill_ratings(self, data: ProspectData) -> None:
        """Add skill ratings table."""
        if not data.skill_ratings:
            return

        self.doc.add_heading("SKILL RATINGS", level=1)

        skills_table = self.doc.add_table(rows=len(data.skill_ratings) + 1, cols=2)
        skills_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        skills_table.rows[0].cells[0].text = "Skill"
        skills_table.rows[0].cells[1].text = "Percentile"
        for cell in skills_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True
            self._set_cell_shading(cell, self.SKILLS_HEADER_COLOR)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for i, (skill, rating) in enumerate(data.skill_ratings.items(), 1):
            skills_table.rows[i].cells[0].text = skill
            skills_table.rows[i].cells[1].text = f"{rating}%"
            skills_table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()

    def _add_recruiting_grades(self, data: ProspectData) -> None:
        """Add recruiting grades section."""
        if not any([data.espn_rating, data.rating_247, data.rivals_rating]):
            return

        self.doc.add_heading("RECRUITING GRADES", level=1)

        grades_para = self.doc.add_paragraph()
        if data.espn_rating:
            grades_para.add_run("ESPN: ").bold = True
            grades_para.add_run(f"{data.espn_rating}   ")
        if data.rating_247:
            grades_para.add_run("247 Sports: ").bold = True
            grades_para.add_run(f"{data.rating_247}   ")
        if data.rivals_rating:
            grades_para.add_run("Rivals: ").bold = True
            grades_para.add_run(data.rivals_rating)

        self.doc.add_paragraph()

    def _add_comparisons(self, data: ProspectData) -> None:
        """Add player comparisons table."""
        if not data.comparisons:
            return

        self.doc.add_heading("PLAYER COMPARISONS", level=1)

        comp_table = self.doc.add_table(rows=len(data.comparisons) + 1, cols=3)
        comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        comp_headers = ["Player", "School", "Similarity"]
        for i, header in enumerate(comp_headers):
            cell = comp_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            self._set_cell_shading(cell, self.COMPARISONS_HEADER_COLOR)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for i, (name, school, similarity) in enumerate(data.comparisons, 1):
            comp_table.rows[i].cells[0].text = name
            comp_table.rows[i].cells[1].text = school
            comp_table.rows[i].cells[2].text = f"{similarity}%"
            comp_table.rows[i].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()

    def _add_bio(self, data: ProspectData) -> None:
        """Add biography section."""
        if not data.bio:
            return

        self.doc.add_heading("BACKGROUND", level=1)
        self._add_text_paragraphs(data.bio)
        self.doc.add_paragraph()

    def _add_strengths(self, data: ProspectData) -> None:
        """Add strengths section."""
        if not data.strengths:
            return

        self.doc.add_heading("SCOUTING REPORT: STRENGTHS", level=1)
        for strength in data.strengths:
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(strength)
        self.doc.add_paragraph()

    def _add_weaknesses(self, data: ProspectData) -> None:
        """Add weaknesses section."""
        if not data.weaknesses:
            return

        self.doc.add_heading("SCOUTING REPORT: WEAKNESSES", level=1)
        for weakness in data.weaknesses:
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(weakness)
        self.doc.add_paragraph()

    def _add_summary(self, data: ProspectData) -> None:
        """Add scouting summary section."""
        if not data.summary:
            return

        self.doc.add_heading("SCOUTING SUMMARY", level=1)
        self._add_text_paragraphs(data.summary)
        self.doc.add_paragraph()

    def _add_consensus_rankings(self, data: ProspectData) -> None:
        """Add consensus rankings section."""
        if not data.avg_overall_rank and not data.avg_position_rank:
            return

        self.doc.add_heading("CONSENSUS RANKINGS", level=1)

        consensus_para = self.doc.add_paragraph()
        if data.avg_overall_rank:
            consensus_para.add_run("Average Overall Rank: ").bold = True
            consensus_para.add_run(f"#{data.avg_overall_rank}   ")
        if data.avg_position_rank:
            consensus_para.add_run("Average Position Rank: ").bold = True
            consensus_para.add_run(f"#{data.avg_position_rank}")

    def _add_footer(self, data: ProspectData) -> None:
        """Add document footer."""
        self.doc.add_paragraph()
        footer = self.doc.add_paragraph("─" * 60)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        source_para = self.doc.add_paragraph()
        source_para.add_run("Source: NFLDraftBuzz.com").italic = True
        source_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if data.last_updated:
            updated_para = self.doc.add_paragraph()
            updated_para.add_run(f"Last Updated: {data.last_updated}").italic = True
            updated_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_text_paragraphs(self, text: str) -> None:
        """Add text split into paragraphs."""
        paragraphs = text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                p = self.doc.add_paragraph(para_text.strip())
                p.paragraph_format.space_after = Pt(12)

    @staticmethod
    def _set_cell_shading(cell, color: str) -> None:
        """Set background color for a table cell."""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)


class DraftBuzzScraper:
    """Main orchestrator for scraping NFL Draft Buzz prospect pages."""

    def __init__(self,
                 playwright: Playwright,
                 profile_root_dir: Path = None,
                 fetcher: PageFetcher = None,
                 doc_generator: WordDocumentGenerator = None):
        self.profile_root_dir = profile_root_dir
        self.base_url = "https://www.nfldraftbuzz.com"
        self.fetcher = fetcher or PageFetcher(playwright=playwright,
                                              base_url=self.base_url)
        self.parser = None
        self.doc_generator = doc_generator or WordDocumentGenerator()
        self.position_rankings_used = defaultdict(list)

    def scrape_from_url(self, url: str, position: str) -> ProspectDataSoup:
        """Scrape prospect data from a URL."""
        print("Parsing prospect data...")
        full_url = f"{self.base_url}{url}"
        base_soup = self.fetcher.fetch_soup(url=full_url)
        self.parser = ProspectParserSoup(soup=base_soup,
                                         position=position)
        prospect_data = self.parser.parse()

        print("Fetching stats page")
        slug_parts = url.split("/")
        player_stats_slug = f"/{slug_parts[1]}/stats/{slug_parts[-1]}"
        stats_full_url = f"{self.base_url}{player_stats_slug}"

        stats_soup = self.fetcher.fetch_soup(url=stats_full_url)
        print("Attempting to parse stats")
        stats_data = self.parser.parse_stats(soup=stats_soup)
        prospect_data.stats = stats_data

        return prospect_data

    def generate_document(self, data: ProspectData, output_path: str) -> None:
        """Generate a Word document from prospect data."""
        print(f"Creating document for: {data.name}")
        print(f"Filename: {output_path}")
        self.doc_generator.generate(data, output_path)

    def _get_prospect_position_rank(self, prospect_data: ProspectData):
        current_rank = int(prospect_data.position_rank)
        rank_is_valid = False
        while not rank_is_valid:
            if current_rank not in self.position_rankings_used[prospect_data.position]:
                rank_is_valid = True
                self.position_rankings_used[prospect_data.position].append(current_rank)
            else:
                print(f"Attempting to use rank that has already been assigned: {current_rank}")
                print("Searching for new one")
                current_rank += 1

        return str(current_rank)

    def generate_output_path(self, data: ProspectData) -> str:
        """Generate default output path from prospect name."""
        safe_name = re.sub(r'[^\w\s-]', '', data.name).replace(' ', '_')
        rank = self._get_prospect_position_rank(prospect_data=data)
        padded_pos_rank = rank.zfill(2)
        return f"{padded_pos_rank}_{safe_name}_Scouting_Report.docx"

    def print_summary(self, data: ProspectData) -> None:
        """Print summary of extracted data."""
        print("\nExtracted data summary:")
        print(f"  Name: {data.name}")
        print(f"  Position: {data.position}")
        print(f"  School: {data.school}")
        print(f"  Rating: {data.overall_rating}/100")
        print(f"  Draft Projection: {data.draft_projection}")
        print(f"  Strengths: {len(data.strengths)} items")
        print(f"  Weaknesses: {len(data.weaknesses)} items")
        print(f"  Image: {'Yes' if data.image_data else 'No'}")


class ProspectProfileListExtractor:
    MAX_RETRIES = 3

    def __init__(self, playwright: Playwright):
        self.playwright = playwright
        self.browser = self._launch_browser()
        self.base_url = "https://www.nfldraftbuzz.com"

    def _launch_browser(self) -> Browser:
        """Launch a new browser instance."""
        return self.playwright.firefox.launch(headless=False)

    def _ensure_browser_connected(self) -> None:
        """Ensure browser is connected, relaunch if necessary."""
        if not self.browser.is_connected():
            print("Browser disconnected, relaunching...")
            self.browser = self._launch_browser()

    def _navigate_with_retry(self, page, url: str) -> None:
        """Navigate to URL with retry logic for browser crashes."""
        for attempt in range(self.MAX_RETRIES):
            try:
                page.goto(url)
                return
            except PlaywrightError as e:
                error_msg = str(e).lower()
                if "target closed" in error_msg or "browser has been closed" in error_msg:
                    print(f"Browser/target closed during navigation (attempt {attempt + 1}/{self.MAX_RETRIES})")
                    raise  # Let caller handle browser relaunch
                raise

    def extract_prospect_hrefs(self, page):
        print(f"Extracting prospect hrefs for {page.url}")
        rows = page.locator('#positionRankTable tbody tr')
        data_hrefs = rows.evaluate_all("rows => rows.map(row => row.getAttribute('data-href'))")
        return data_hrefs

    def extract_prospect_urls_for_position(self, pos: str) -> List[str]:
        all_profiles = []

        path = f"/positions/{pos}/1/2026"
        full_url = f"{self.base_url}{path}"

        page = self._create_page_with_retry(full_url)
        all_profiles.extend(self.extract_prospect_hrefs(page=page))
        links = page.locator('ul.pagination li.page-item a.page-link[href]')
        position_page_hrefs = links.evaluate_all("anchors => anchors.map(a => a.getAttribute('href'))")

        for path in position_page_hrefs:
            page.close()
            full_url = f"{self.base_url}{path}"
            page = self._create_page_with_retry(full_url)
            time.sleep(uniform(4.5, 5.5))

            prospect_hrefs = self.extract_prospect_hrefs(page)
            all_profiles.extend(prospect_hrefs)
        page.close()
        return all_profiles

    def _create_page_with_retry(self, url: str):
        """Create a new page and navigate to URL with retry on browser crash."""
        last_error = None
        for attempt in range(self.MAX_RETRIES):
            try:
                self._ensure_browser_connected()
                page = self.browser.new_page()
                page.goto(url, timeout=0)
                return page
            except PlaywrightError as e:
                last_error = e
                error_msg = str(e).lower()
                if "target closed" in error_msg or "browser has been closed" in error_msg:
                    print(f"Browser/target closed (attempt {attempt + 1}/{self.MAX_RETRIES}), relaunching...")
                    try:
                        self.browser.close()
                    except Exception:
                        pass
                    self.browser = self._launch_browser()
                    time.sleep(1)
                else:
                    raise
        raise last_error
